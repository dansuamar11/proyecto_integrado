# -*- coding: utf-8 -*-
from docx import Document
path = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"
doc = Document(path)
for i in range(303, min(len(doc.paragraphs), 312)):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f"{i}: [{p.style.name}] {t}")
