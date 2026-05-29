# -*- coding: utf-8 -*-
from docx import Document
path = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"
doc = Document(path)
indices = [110,141,191,230,249,253,279]
for idx in indices:
    print(f"\n--- around {idx} ---")
    for i in range(max(0, idx-2), min(len(doc.paragraphs), idx+12)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if text:
            print(f"{i}: [{p.style.name}] {text}")
