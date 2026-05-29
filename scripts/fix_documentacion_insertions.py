# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor

DOC_PATH = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"


def contains_text(document, text):
    return any(text in paragraph.text for paragraph in document.paragraphs)


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(text)


def insert_after(paragraph, text, style_name, red=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph.style = style_name
    run = new_paragraph.add_run(text)
    if red:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return new_paragraph


doc = Document(DOC_PATH)

backend_text = "En esa logica de negocio se han concentrado tambien las validaciones de generacion de partidos: restriccion de franjas horarias, deteccion de colisiones exactas de fecha y hora, y comprobacion de que un mismo equipo no quede programado en dos partidos distintos durante el mismo dia."
if not contains_text(doc, backend_text):
    anchor = find_paragraph(doc, "Esta separaci")
    insert_after(anchor, backend_text, "Normal")

test_anchor = "Comprobación del cambio de contraseña para perfiles autenticados y de la redireccion a 404 al intentar abrir rutas privadas sin sesion."
test_entries = [
    "Comprobación de la generación de partidos con franjas horarias cerradas, bloqueo de colisiones de fecha y hora y bloqueo de programación doble de un mismo equipo en el mismo día.",
    "Comprobación de la reasignación de árbitros en partidos pendientes y de la visualización del aviso flotante con cierre manual sobre el modal.",
    "Comprobación del comando raíz npm run dev para verificar el arranque coordinado de backend y frontend.",
]

anchor = find_paragraph(doc, test_anchor)
last = anchor
for entry in test_entries:
    if contains_text(doc, entry):
        continue
    last = insert_after(last, entry, "List Bullet")

doc.save(DOC_PATH)
