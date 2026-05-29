# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"

BACKEND_TEXT = "En esa logica de negocio se han concentrado tambien las validaciones de generacion de partidos: restriccion de franjas horarias, deteccion de colisiones exactas de fecha y hora, y comprobacion de que un mismo equipo no quede programado en dos partidos distintos durante el mismo dia."
TEST_TEXTS = [
    "Comprobación de la generación de partidos con franjas horarias cerradas, bloqueo de colisiones de fecha y hora y bloqueo de programación doble de un mismo equipo en el mismo día.",
    "Comprobación de la reasignación de árbitros en partidos pendientes y de la visualización del aviso flotante con cierre manual sobre el modal.",
    "Comprobación del comando raíz npm run dev para verificar el arranque coordinado de backend y frontend.",
]


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(text)


def find_all_paragraphs(document, text):
    return [paragraph for paragraph in document.paragraphs if text in paragraph.text]


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def insert_after(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph.style = style_name
    new_paragraph.add_run(text)
    return new_paragraph


doc = Document(DOC_PATH)

for paragraph in find_all_paragraphs(doc, BACKEND_TEXT):
    remove_paragraph(paragraph)

for text in TEST_TEXTS:
    for paragraph in find_all_paragraphs(doc, text):
        remove_paragraph(paragraph)

backend_anchor = find_paragraph(doc, "Esta separaci")
insert_after(backend_anchor, BACKEND_TEXT, "Normal")

tests_anchor = find_paragraph(doc, "Comprobación del cambio de contraseña para perfiles autenticados")
last = tests_anchor
for text in TEST_TEXTS:
    last = insert_after(last, text, "List Bullet")

doc.save(DOC_PATH)
