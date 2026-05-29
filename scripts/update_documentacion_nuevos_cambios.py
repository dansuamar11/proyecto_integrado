# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

DOC_PATH = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"


def insert_paragraph_after(paragraph, text, style_name=None, red=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if style_name:
        new_paragraph.style = style_name
    run = new_paragraph.add_run(text)
    if red:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return new_paragraph


def contains_text(document, text):
    return any(text in paragraph.text for paragraph in document.paragraphs)


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"No se encontro el texto ancla: {text}")


doc = Document(DOC_PATH)

updates = [
    (
        "Del mismo modo, la reasignacion de entrenador de un equipo se actualiza ya de forma inmediata",
        [
            ("Tambien se ha incorporado la generacion de partidos desde un modal especifico del panel de gerente. Este flujo restringe las horas disponibles a las 12:00, 13:00, 16:00, 17:00 y 18:00, valida que no exista otro partido ya programado en esa misma fecha y franja y evita que un mismo equipo pueda disputar dos encuentros el mismo dia.", "Normal", False),
            ("Ademas, la card de asignacion de arbitros ya no se limita a partidos sin arbitro, sino que permite asignar o reasignar el arbitro de cualquier partido pendiente. Esto resulta util cuando un arbitro causa baja o es necesario redistribuir designaciones sin modificar el resto de la informacion del encuentro.", "Normal", False),
            ("AÑADIR CAPTURA DEL MODAL DE GENERAR PARTIDO CON EL AVISO SUPERPUESTO DE VALIDACION", "Normal", True),
        ],
    ),
    (
        "Una vez finalizado el arranque",
        [
            ("Para desarrollo local sin Docker completo, el proyecto incorpora tambien un comando raiz \"npm run dev\" que levanta de forma coordinada el backend Spring Boot y el frontend Angular, esperando a que ambos servicios queden realmente disponibles antes de dar el entorno por iniciado.", "Normal", False),
        ],
    ),
    (
        "Si al dar de alta un nuevo usuario introduce un username ya existente",
        [
            ("Asimismo, desde el panel puede generar partidos nuevos indicando equipos, fecha y arbitro opcional. La hora se selecciona de una lista cerrada de franjas validas, y si se intenta reutilizar una fecha y hora ya ocupadas o programar a un equipo dos veces el mismo dia, la aplicacion muestra un aviso superpuesto sobre el propio modal y bloquea la creacion.", "List Paragraph", False),
            ("El gerente tambien puede asignar y reasignar arbitros sobre cualquier partido pendiente, de modo que una baja de arbitro no obliga a recrear el encuentro ni a tocar manualmente la base de datos.", "List Paragraph", False),
            ("AÑADIR CAPTURA DEL MODAL DE PARTIDOS MOSTRANDO EL BOTON DE CIERRE DEL AVISO", "Normal", True),
        ],
    ),
    (
        "A nivel de comportamiento de interfaz, se han ajustado las actualizaciones inmediatas del panel de gerente",
        [
            ("Sobre esa misma base se ha incorporado un aviso flotante y cerrable para los errores de generacion de partidos, mostrado por encima del modal cuando existe una colision de fecha y hora o cuando alguno de los equipos ya tiene otro partido ese mismo dia.", "Normal", False),
        ],
    ),
    (
        "Esta separaci",
        [
            ("En esa logica de negocio se han concentrado tambien las validaciones de generacion de partidos: restriccion de franjas horarias, deteccion de colisiones exactas de fecha y hora, y comprobacion de que un mismo equipo no quede programado en dos partidos distintos durante el mismo dia.", "Normal", False),
        ],
    ),
    (
        "Desde el punto de vista del mantenimiento, este tipo de scripts tambi",
        [
            ("Ademas, el arranque tecnico garantiza ahora la existencia de un usuario administrador inicial con credenciales admin/admin123 cuando la base se levanta por primera vez sin datos previos, facilitando el acceso inicial a las tareas de configuracion y verificacion.", "Normal", False),
        ],
    ),
    (
        "Comprobaci",
        [
            ("Comprobación de la generación de partidos con franjas horarias cerradas, bloqueo de colisiones de fecha y hora y bloqueo de programación doble de un mismo equipo en el mismo día.", "List Bullet", False),
            ("Comprobación de la reasignación de árbitros en partidos pendientes y de la visualización del aviso flotante con cierre manual sobre el modal.", "List Bullet", False),
            ("Comprobación del comando raíz npm run dev para verificar el arranque coordinado de backend y frontend.", "List Bullet", False),
        ],
    ),
]

for anchor_text, paragraphs in updates:
    last_inserted = find_paragraph(doc, anchor_text)
    for text, style_name, red in paragraphs:
        if contains_text(doc, text):
            continue
        last_inserted = insert_paragraph_after(last_inserted, text, style_name, red)

doc.save(DOC_PATH)
