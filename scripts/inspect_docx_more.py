# -*- coding: utf-8 -*-
from docx import Document
path = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"
doc = Document(path)
for start in [200, 260, 286]:
    print(f"\n--- block {start} ---")
    for i in range(start, min(len(doc.paragraphs), start+18)):
        p = doc.paragraphs[i]
        t = p.text.strip()
        if t:
            print(f"{i}: [{p.style.name}] {t}")
