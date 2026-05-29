from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def find_paragraph(paragraphs, startswith_text):
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if text.startswith(startswith_text):
            return paragraph
    raise ValueError(f"No se encontro el parrafo que empieza por: {startswith_text!r}")


def add_red_note(paragraph, text):
    note = insert_paragraph_after(paragraph, style=paragraph._parent.part.document.styles["Normal"])
    run = note.add_run(text.upper())
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    return note


def main():
    doc_path = Path(r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx")
    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    # 6.2 Sistema de autenticacion y control de acceso
    p = find_paragraph(paragraphs, "Este control de acceso diferencia")
    p = insert_paragraph_after(
        p,
        "Como refuerzo reciente de esta capa, las rutas privadas del frontend redirigen ahora a la pagina 404 cuando se intenta acceder sin haber iniciado sesion o sin disponer del rol adecuado. De esta forma se evita cargar vistas internas con informacion incompleta o mensajes de error funcionales cuando el acceso no corresponde al perfil autenticado.",
        style=doc.styles["Normal"],
    )
    p = insert_paragraph_after(
        p,
        "Tambien se ha incorporado una funcionalidad de cambio de contraseña para los perfiles autenticados, accesible desde sus respectivos paneles. El proceso exige introducir la contraseña actual y escribir dos veces la nueva, manteniendo una validacion tanto en frontend como en backend para asegurar la coherencia del cambio.",
        style=doc.styles["Normal"],
    )

    # 6.3 Gestion de gerente
    p = find_paragraph(paragraphs, "Además, el rol de gerente incluye una página")
    p = insert_paragraph_after(
        p,
        "En la gestion de usuarios tambien se ha añadido una validacion especifica para impedir altas con un nombre de usuario ya existente en base de datos. Cuando se intenta reutilizar un username registrado, el sistema muestra un aviso en rojo y cancela la operacion sin crear el nuevo perfil.",
        style=doc.styles["Normal"],
    )
    p = insert_paragraph_after(
        p,
        "Del mismo modo, la reasignacion de entrenador de un equipo se actualiza ya de forma inmediata en la interfaz, mostrando el aviso de confirmacion en el primer intento y refrescando al momento la informacion visible en el panel.",
        style=doc.styles["Normal"],
    )
    add_red_note(p, "AÑADIR CAPTURA DEL ALERT DE USERNAME DUPLICADO EN EL PANEL DE GERENTE")

    # 6.4 Gestion de entrenador
    p = find_paragraph(paragraphs, "La vista de entrenador está orientada")
    p = insert_paragraph_after(
        p,
        "Tambien se ha contemplado el caso en el que un entrenador todavia no tenga equipo asignado. En esa situacion, el panel muestra un mensaje explicito indicando que no dispone aun de equipo, evitando errores de carga o vistas vacias ambiguas.",
        style=doc.styles["Normal"],
    )
    p = insert_paragraph_after(
        p,
        "Respecto a la relacion con las actas arbitrales, el estado activo o inactivo de los jugadores adquiere relevancia directa: los jugadores dados de baja dejan de aparecer en las fichas abiertas de los partidos futuros o pendientes, aunque se mantienen en el historial de encuentros ya jugados si registraron estadisticas en ellos.",
        style=doc.styles["Normal"],
    )
    add_red_note(p, 'AÑADIR CAPTURA DEL MENSAJE "NO TIENES EQUIPO ASIGNADO AÚN" EN EL PANEL DE ENTRENADOR')

    # 6.5 Gestion de arbitro y actas
    p = find_paragraph(paragraphs, "Este bloque funcional resulta especialmente importante")
    p = insert_paragraph_after(
        p,
        "En esta parte del proyecto se ha refinado la construccion de las actas para distinguir entre partidos abiertos y partidos ya jugados. En los partidos pendientes solo se muestran jugadores activos de los equipos implicados, mientras que en los encuentros ya disputados se preserva la informacion historica de quienes participaron, aunque despues hayan pasado a estar inactivos.",
        style=doc.styles["Normal"],
    )
    p = insert_paragraph_after(
        p,
        "Ademas, el arbitro dispone ahora del mismo mecanismo de cambio de contraseña integrado en su panel, siguiendo la misma validacion de contraseña actual, doble confirmacion de la nueva y bloqueo de cambios inconsistentes.",
        style=doc.styles["Normal"],
    )
    add_red_note(p, "AÑADIR CAPTURA DEL MODAL DE CAMBIO DE CONTRASEÑA O DE UNA FICHA ABIERTA/JUGADA")

    # 9.2 Inicio de sesion
    p = find_paragraph(paragraphs, "Este comportamiento permite que cada perfil vea")
    insert_paragraph_after(
        p,
        "Si un usuario intenta acceder manualmente a una ruta privada que no le corresponde o lo hace sin sesion iniciada, la aplicacion redirige a la pagina 404 en lugar de cargar la vista privada con informacion capada. Este comportamiento mejora la coherencia de navegacion y deja mas claro que la ruta no esta disponible para ese contexto.",
        style=doc.styles["Normal"],
    )

    # 9.3 Gerente
    p = find_paragraph(paragraphs, "Gestionar las solicitudes de contacto mediante")
    p = insert_paragraph_after(
        p,
        "Tambien puede cambiar su propia contraseña desde un modal accesible en el panel, introduciendo la contraseña actual y repitiendo dos veces la nueva.",
        style=doc.styles["List Paragraph"],
    )
    insert_paragraph_after(
        p,
        "Si al dar de alta un nuevo usuario introduce un username ya existente, la aplicacion muestra un aviso en rojo indicando que ese nombre de usuario ya esta en uso y obliga a escoger otro distinto antes de completar el alta.",
        style=doc.styles["List Paragraph"],
    )

    # 9.4 Entrenador
    p = find_paragraph(paragraphs, "La idea de este perfil es que el entrenador disponga")
    p = insert_paragraph_after(
        p,
        "Si todavia no tiene un equipo asignado, el sistema se lo comunica de forma explicita mediante el mensaje 'No tienes equipo asignado aún', evitando cargar un panel sin contexto.",
        style=doc.styles["Normal"],
    )
    p = insert_paragraph_after(
        p,
        "El entrenador tambien puede cambiar su contraseña desde su panel. Ademas, el estado activo o inactivo de sus jugadores afecta a las futuras actas arbitrales: los jugadores de baja no apareceran en fichas abiertas, aunque seguiran figurando en partidos ya jugados si llegaron a participar en ellos.",
        style=doc.styles["Normal"],
    )

    # 9.5 Arbitro
    p = find_paragraph(paragraphs, "Cada una de estas cards cuenta con un botón")
    p = insert_paragraph_after(
        p,
        "En las fichas de partidos aun abiertos, el arbitro solo visualiza jugadores activos de cada equipo. Sin embargo, una vez que el partido ya ha sido jugado, la ficha conserva a los jugadores que registraron estadisticas, incluso si mas adelante se les da de baja por parte del entrenador.",
        style=doc.styles["Normal"],
    )
    insert_paragraph_after(
        p,
        "Como en el resto de perfiles privados, el arbitro dispone tambien de un modal para cambiar su contraseña desde su propia vista.",
        style=doc.styles["Normal"],
    )

    # 10.2 Frontend
    p = find_paragraph(paragraphs, "Desde un punto de vista técnico, Angular se apoya")
    p = insert_paragraph_after(
        p,
        "Dentro del frontend se han añadido tambien guardas de ruta mas estrictas para que los accesos privados sin sesion o sin rol valido redirijan a la pagina 404, asi como componentes reutilizables para el cambio de contraseña mediante modal en los distintos paneles privados.",
        style=doc.styles["Normal"],
    )
    insert_paragraph_after(
        p,
        "A nivel de comportamiento de interfaz, se han ajustado las actualizaciones inmediatas del panel de gerente para que avisos como el de username duplicado o el cambio de entrenador se reflejen en pantalla en el momento en que llega la respuesta del servidor.",
        style=doc.styles["Normal"],
    )

    # 10.3 Backend
    p = find_paragraph(paragraphs, "En esta parte del backend se han añadido endpoints")
    p = insert_paragraph_after(
        p,
        "Sobre esta base se ha ampliado tambien la API privada con un endpoint especifico de cambio de contraseña y con validaciones adicionales, como la comprobacion de usernames duplicados al crear usuarios internos o la construccion diferenciada de actas segun el estado del partido.",
        style=doc.styles["Normal"],
    )
    insert_paragraph_after(
        p,
        "Asimismo, el backend contempla ya el caso del entrenador sin equipo asignado, devolviendo una respuesta controlada para que el frontend pueda mostrar un mensaje informativo en lugar de un error funcional.",
        style=doc.styles["Normal"],
    )

    # 10.4 Seguridad
    p = find_paragraph(paragraphs, "Además, el backend protege sus endpoints privados")
    p = insert_paragraph_after(
        p,
        "La seguridad aplicada no se limita al control de acceso inicial: el cambio de contraseña exige validar la contraseña actual y confirmar dos veces la nueva, mientras que la creacion de usuarios internos incorpora restricciones de integridad para impedir usernames duplicados.",
        style=doc.styles["Normal"],
    )

    # 11 Pruebas
    p = find_paragraph(paragraphs, "Comprobación del despliegue general mediante Docker Compose.")
    p = insert_paragraph_after(
        p,
        "Comprobación del bloqueo de alta de usuarios con username duplicado y visualizacion inmediata del aviso en el panel de gerente.",
        style=doc.styles["List Bullet"],
    )
    p = insert_paragraph_after(
        p,
        "Comprobación del mensaje mostrado al entrenador sin equipo asignado y del comportamiento de jugadores activos e inactivos en fichas abiertas y jugadas.",
        style=doc.styles["List Bullet"],
    )
    insert_paragraph_after(
        p,
        "Comprobación del cambio de contraseña para perfiles autenticados y de la redireccion a 404 al intentar abrir rutas privadas sin sesion.",
        style=doc.styles["List Bullet"],
    )

    doc.save(doc_path)


if __name__ == "__main__":
    main()
