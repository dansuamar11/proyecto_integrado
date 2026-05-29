# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml import OxmlElement

DOC_PATH = r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx"
TEXT = "En esa logica de negocio se han concentrado tambien las validaciones de generacion de partidos: restriccion de franjas horarias, deteccion de colisiones exactas de fecha y hora, y comprobacion de que un mismo equipo no quede programado en dos partidos distintos durante el mismo dia."


def remove_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def insert_after(paragraph, text, style_name):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph.style = style_name
    new_paragraph.add_run(text)
    return new_paragraph


doc = Document(DOC_PATH)

for paragraph in list(doc.paragraphs):
    if TEXT in paragraph.text:
        remove_paragraph(paragraph)

anchor = None
for paragraph in doc.paragraphs:
    if "El backend ha sido desarrollado con Spring Boot" in paragraph.text:
        anchor = paragraph
        break

if anchor is None:
    raise ValueError("No se encontro el bloque del backend")

insert_after(anchor, TEXT, "Normal")
doc.save(DOC_PATH)
