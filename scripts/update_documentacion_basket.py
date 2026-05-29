from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


TARGET_DOC = Path(r"C:\Users\DSuar\OneDrive\Escritorio\TFG\Documentacion basket aljarafe.docx")


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def find_paragraph(document: Document, needle: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == needle:
            return paragraph
    raise ValueError(f"No se ha encontrado el apartado exacto: {needle}")


def clear_section(document: Document, start_heading: str, end_heading: str) -> Paragraph:
    start_paragraph = find_paragraph(document, start_heading)
    end_paragraph = find_paragraph(document, end_heading)
    delete_list: list[Paragraph] = []
    cursor = start_paragraph._element.getnext()
    while cursor is not None and cursor is not end_paragraph._element:
        for paragraph in document.paragraphs:
            if paragraph._element is cursor:
                delete_list.append(paragraph)
                break
        cursor = cursor.getnext()

    for paragraph in delete_list:
        remove_paragraph(paragraph)

    return start_paragraph


def append_block(anchor: Paragraph, lines: Iterable[tuple[str, str | None]]) -> Paragraph:
    current = anchor
    for text, style in lines:
        current = insert_paragraph_after(current, text, style)
    return current


def build_contact_public_block() -> list[tuple[str, str | None]]:
    return [
        (
            "Dentro de la zona publica se ha incorporado tambien una pagina especifica de contacto, pensada para que cualquier persona interesada pueda comunicarse con el club sin necesidad de disponer de una cuenta.",
            None,
        ),
        (
            "En esta pantalla el usuario puede rellenar un formulario con su nombre, correo electronico, telefono y mensaje. El sistema obliga a indicar el mensaje y, ademas, a facilitar al menos una via de respuesta, ya sea el correo o el telefono, de forma que la solicitud siempre llegue con informacion minima util para su posterior gestion.",
            None,
        ),
        (
            "Una vez enviado el formulario, la informacion se almacena en la aplicacion y pasa a estar disponible para el gerente desde la parte privada. De esta manera, la pagina de contacto no se queda solo en una seccion informativa, sino que funciona como un canal real de comunicacion entre la web y la administracion del club.",
            None,
        ),
        ("INTRODUCIR CAPTURAS DE LA PAGINA PUBLICA DE CONTACTO", None),
        ("INTRODUCIR CAPTURAS DEL FORMULARIO DE CONTACTO RELLENO Y DEL MENSAJE DE ENVIO CORRECTO", None),
    ]


def build_contact_admin_block() -> list[tuple[str, str | None]]:
    return [
        (
            "Ademas de las funciones habituales de administracion, el gerente dispone ahora de una seccion propia para revisar las solicitudes recibidas desde la pagina publica de contacto.",
            None,
        ),
        (
            "Desde el panel principal del gerente se muestra un acceso directo a esta funcionalidad junto con un indicador del numero de solicitudes pendientes. Al entrar, el gerente puede consultar de forma ordenada la fecha de llegada, el nombre de la persona, su telefono, su correo electronico y el contenido completo del mensaje enviado.",
            None,
        ),
        (
            "El funcionamiento previsto es sencillo: el gerente entra en la seccion, revisa las solicitudes pendientes, contacta con la persona por la via que considere oportuna y, cuando la solicitud ya ha sido gestionada, la elimina desde el propio panel para mantener el listado limpio y actualizado.",
            None,
        ),
        (
            "Esta parte resulta util porque centraliza la comunicacion externa en un unico punto del sistema y evita que los mensajes queden dispersos o se pierdan, permitiendo que el club lleve un control mas comodo de las peticiones que llegan desde la web.",
            None,
        ),
        ("INTRODUCIR CAPTURAS DEL ACCESO A SOLICITUDES DE CONTACTO DESDE EL PANEL DE GERENTE", None),
        ("INTRODUCIR CAPTURAS DEL LISTADO DE SOLICITUDES DE CONTACTO Y DEL BOTON DE BORRADO", None),
    ]


def build_section_8() -> list[tuple[str, str | None]]:
    return [
        ("En este apartado se explica de forma mas detallada como poner en funcionamiento el proyecto, tanto en un equipo local como en un servidor. La idea del despliegue es que la aplicacion pueda arrancarse de manera comoda mediante contenedores Docker, separando la base de datos MySQL, el backend desarrollado con Spring Boot y el frontend desarrollado con Angular.", None),
        ("8.1 Requisitos previos", "Heading 2"),
        ("Para realizar la instalacion del proyecto es necesario disponer, como minimo, de Docker y Docker Compose correctamente instalados en el equipo o servidor donde se vaya a desplegar la aplicacion. En caso de trabajar en local para pruebas o demostraciones, tambien es recomendable contar con un editor de codigo y una terminal desde la que poder ejecutar los comandos necesarios.", None),
        ("Ademas, conviene comprobar que los puertos que utiliza el sistema se encuentran disponibles. En la configuracion actual se emplea el puerto 3307 para MySQL, el 1001 para el backend y el 80 para el frontend, aunque estos valores pueden modificarse facilmente mediante variables de entorno.", None),
        ("8.2 Ficheros necesarios para el despliegue", "Heading 2"),
        ("El proyecto incluye un fichero compose.yaml desde el que se orquesta todo el entorno. En dicho fichero se definen los tres servicios principales del sistema: la base de datos MySQL, el backend y el frontend. Tambien se utiliza un fichero .env para centralizar las variables de configuracion mas importantes, como el nombre de la base de datos, usuarios, contrasenas, puertos y zona horaria.", None),
        ("En la parte de base de datos, ademas, el proyecto cuenta con scripts de inicializacion y carga de datos de ejemplo, de manera que la primera puesta en marcha del sistema resulte mucho mas comoda si se desea partir de una base ya preparada para pruebas o demostraciones.", None),
        ("INTRODUCIR CAPTURAS DE LOS FICHEROS DE DESPLIEGUE DEL PROYECTO", None),
        ("8.3 Puesta en marcha desde el codigo fuente", "Heading 2"),
        ("Si se quiere desplegar el sistema directamente a partir del codigo fuente, el procedimiento general consiste en situarse en la carpeta raiz del proyecto, revisar el fichero .env para adaptar los valores al entorno de trabajo y, a continuacion, ejecutar el comando docker compose up -d --build.", None),
        ("Con esta orden Docker construye la imagen del backend a partir del Dockerfile del proyecto Java, construye tambien la imagen del frontend a partir del Dockerfile ubicado en la aplicacion Angular y levanta junto a ellas el contenedor de MySQL. El uso de la opcion --build es especialmente util cuando se han realizado cambios recientes en el codigo y se desea asegurar que las imagenes se generen de nuevo con la version actual del proyecto.", None),
        ("Una vez finalizado el arranque, se puede comprobar el estado de los contenedores con docker compose ps o docker ps. Si todo ha ido correctamente, el frontend deberia estar accesible desde el navegador y el backend deberia encontrarse operativo recibiendo las peticiones de la aplicacion cliente.", None),
        ("INTRODUCIR CAPTURAS DEL LEVANTAMIENTO DEL PROYECTO CON DOCKER COMPOSE DESDE EL CODIGO FUENTE", None),
        ("8.4 Despliegue en un servidor utilizando las imagenes publicadas en Docker Hub", "Heading 2"),
        ("Una de las ventajas de la configuracion actual es que el proyecto tambien puede desplegarse sin necesidad de compilarlo en el propio servidor, siempre que las imagenes del frontend y del backend se encuentren publicadas previamente en Docker Hub. Esto resulta especialmente util cuando se quiere hacer una puesta en produccion o una demostracion en un servidor remoto de una forma mas rapida y controlada.", None),
        ("En ese caso, en lugar de utilizar bloques build dentro del compose, bastaria con sustituirlos por referencias image apuntando a las imagenes publicadas. Por ejemplo, el backend podria tomar una imagen del tipo TU_USUARIO_DOCKERHUB/basket-aljarafe-backend:TAG y el frontend una imagen del tipo TU_USUARIO_DOCKERHUB/basket-aljarafe-frontend:TAG. Una vez ajustado el compose, el servidor solo tendria que descargar dichas imagenes y levantar los servicios.", None),
        ("El flujo habitual en un servidor seria el siguiente: copiar el compose y el fichero .env al servidor, ejecutar docker compose pull para descargar las imagenes mas recientes publicadas en Docker Hub y, seguidamente, lanzar docker compose up -d. Con ello se consigue desplegar la aplicacion de manera bastante limpia, dejando el servidor como mero entorno de ejecucion y evitando tener que instalar en el mismo herramientas de compilacion de Java o Angular.", None),
        ("Tambien se puede utilizar este procedimiento para actualizar el sistema en el futuro. Si se publica una nueva version de las imagenes en Docker Hub, bastaria con volver a ejecutar docker compose pull y docker compose up -d para que el servidor descargue la version actualizada y la ponga en marcha.", None),
        ("INTRODUCIR CAPTURAS DE LAS IMAGENES PUBLICADAS EN DOCKER HUB", None),
        ("INTRODUCIR CAPTURAS DEL DESPLIEGUE EN SERVIDOR DESCARGANDO IMAGENES DESDE DOCKER HUB", None),
        ("8.5 Variables de entorno y configuracion", "Heading 2"),
        ("Las variables de entorno permiten adaptar el comportamiento del sistema sin tocar directamente el codigo. Entre las mas importantes se encuentran el nombre de la base de datos, el usuario y contrasena de MySQL, la contrasena del usuario root, los puertos externos asignados a cada servicio y el valor empleado por Hibernate para la gestion del esquema.", None),
        ("Durante el desarrollo puede resultar comodo trabajar con SPRING_JPA_HIBERNATE_DDL_AUTO configurado en update, ya que facilita la evolucion de la base de datos. Sin embargo, en un entorno mas estable o cercano a produccion conviene revisar cuidadosamente este valor y apoyarse, cuando sea necesario, en scripts de base de datos controlados para evitar cambios imprevistos.", None),
        ("8.6 Comprobaciones tras la instalacion", "Heading 2"),
        ("Despues de levantar el proyecto es conveniente realizar una serie de verificaciones basicas. En primer lugar, debe comprobarse que los tres contenedores se encuentran en ejecucion. En segundo lugar, debe verificarse que la pagina publica carga correctamente en el navegador y que se pueden consultar al menos las secciones principales de la web. En tercer lugar, es recomendable iniciar sesion con alguno de los perfiles disponibles para confirmar que la autenticacion y la parte privada funcionan correctamente.", None),
        ("Tambien es aconsejable comprobar que la conexion entre frontend y backend responde adecuadamente y que la base de datos mantiene los datos entre reinicios gracias al volumen persistente configurado para MySQL.", None),
        ("INTRODUCIR CAPTURAS DE LA COMPROBACION FINAL DEL SISTEMA EN FUNCIONAMIENTO", None),
        ("8.7 Parada, mantenimiento y actualizacion", "Heading 2"),
        ("Para detener temporalmente el sistema se puede utilizar docker compose stop, mientras que para eliminar los contenedores en ejecucion y volver a levantarlos mas adelante se puede recurrir a docker compose down. Si se desea conservar la informacion almacenada en la base de datos, no debe eliminarse el volumen de persistencia asociado a MySQL.", None),
        ("En cuanto al mantenimiento, el procedimiento normal consistiria en actualizar el codigo o publicar nuevas imagenes, volver a construir o descargar las versiones correspondientes y reiniciar los servicios. Si el despliegue se realiza desde Docker Hub, el proceso de mantenimiento se simplifica todavia mas, ya que solamente habria que actualizar las imagenes remotas y volver a hacer pull en el servidor.", None),
    ]


def build_section_9() -> list[tuple[str, str | None]]:
    return [
        ("El manual de usuario describe el funcionamiento general del sistema desde el punto de vista de las personas que van a utilizar la aplicacion. La intencion de este apartado es recoger de forma clara las acciones mas importantes que puede realizar cada tipo de usuario dentro de BASKET_ALJARAFE.", None),
        ("9.1 Acceso a la pagina publica", "Heading 2"),
        ("Cualquier visitante puede acceder a la parte publica de la web sin necesidad de identificarse. Desde esta zona es posible consultar la informacion general del club, navegar por los distintos apartados visibles para el publico y conocer mejor la actividad desarrollada por la entidad.", None),
        ("La parte publica actua como carta de presentacion del proyecto, por lo que esta pensada para ofrecer una navegacion sencilla y directa. El usuario solo tiene que abrir la pagina principal desde el navegador y desplazarse por los distintos apartados usando el menu disponible.", None),
        ("INTRODUCIR CAPTURAS DE LA PAGINA PRINCIPAL Y DE LA NAVEGACION PUBLICA", None),
        ("9.2 Uso de la pagina de contacto", "Heading 2"),
        ("Dentro de la zona publica existe tambien una pagina de contacto destinada a las personas interesadas en comunicarse con el club. El usuario debe rellenar el formulario indicando su nombre si lo desea, una direccion de correo electronico o un numero de telefono, y el mensaje que quiere enviar.", None),
        ("Si el formulario no contiene el mensaje o no se ha indicado ninguna via de contacto valida, la propia aplicacion muestra los avisos oportunos para que el usuario pueda corregir los datos antes de enviarlos. Cuando la informacion es correcta, la solicitud queda registrada y aparece un mensaje confirmando que el envio se ha realizado satisfactoriamente.", None),
        ("Esta funcionalidad esta pensada para facilitar peticiones de informacion, consultas generales o tomas de contacto por parte de familias, jugadores o cualquier otra persona interesada en el club.", None),
        ("INTRODUCIR CAPTURAS DEL PROCESO DE ENVIO DE UNA SOLICITUD DE CONTACTO", None),
        ("9.3 Inicio de sesion", "Heading 2"),
        ("Para acceder a las zonas privadas del sistema es necesario iniciar sesion con un usuario valido. El acceso se realiza desde la pantalla de login, introduciendo las credenciales correspondientes. Una vez autenticado, cada usuario es redirigido al area que le corresponde segun su rol dentro de la aplicacion.", None),
        ("Este comportamiento permite que cada perfil vea solo las herramientas y opciones que realmente necesita, manteniendo separadas las responsabilidades del gerente, del entrenador y del arbitro.", None),
        ("INTRODUCIR CAPTURAS DEL INICIO DE SESION", None),
        ("9.4 Manual de uso para el gerente", "Heading 2"),
        ("El gerente es el perfil con mayores capacidades de gestion dentro del sistema. Desde su panel puede acceder a las opciones necesarias para administrar diferentes elementos de la aplicacion, como la informacion general, los equipos, usuarios u otros recursos internos que se hayan habilitado en el proyecto.", None),
        ("Ademas, el gerente dispone ahora de un acceso especifico a las solicitudes de contacto recibidas desde la pagina publica. Para gestionarlas, debe entrar en su panel principal y pulsar en la opcion correspondiente. Una vez dentro, puede revisar cada solicitud de forma individual, comprobar la fecha en la que llego, leer el mensaje recibido y utilizar el correo o telefono indicados para responder por fuera de la aplicacion si lo considera necesario.", None),
        ("Cuando una solicitud ya ha sido atendida, el gerente puede eliminarla desde el boton de borrado disponible en la misma pantalla. De este modo, el listado refleja un estado mas limpio y se evita mezclar mensajes ya tratados con otros que siguen pendientes.", None),
        ("En conjunto, el uso del panel de gerente busca concentrar en una sola zona las tareas administrativas mas importantes, de manera que la persona responsable del sistema pueda mantener la informacion actualizada con la menor friccion posible.", None),
        ("INTRODUCIR CAPTURAS DEL PANEL PRINCIPAL DE GERENTE", None),
        ("INTRODUCIR CAPTURAS DEL ACCESO Y GESTION DE SOLICITUDES DE CONTACTO POR PARTE DEL GERENTE", None),
        ("9.5 Manual de uso para el entrenador", "Heading 2"),
        ("El entrenador accede a una zona orientada a la gestion deportiva del equipo o de los equipos que tenga asignados. Desde esta parte del sistema puede consultar la informacion que le corresponde, revisar jugadores, realizar las acciones permitidas sobre sus datos y mantener el seguimiento del ambito deportivo dentro de la aplicacion.", None),
        ("La idea de este perfil es que el entrenador disponga de una herramienta concreta para el trabajo diario, sin necesidad de entrar en tareas administrativas que no le corresponden. Por eso la interfaz se centra en mostrar solo aquellas opciones que de verdad le resultan utiles.", None),
        ("INTRODUCIR CAPTURAS DE LA ZONA DE ENTRENADOR Y DE SUS FUNCIONALIDADES PRINCIPALES", None),
        ("9.6 Manual de uso para el arbitro", "Heading 2"),
        ("El arbitro dispone de una zona propia desde la que puede trabajar con los partidos y las actas asociadas. Dependiendo de lo implementado en el proyecto, desde aqui puede consultar encuentros, rellenar o actualizar informacion relacionada con las actas y dejar registrada la informacion arbitral necesaria para el desarrollo del sistema.", None),
        ("Este perfil permite separar la informacion tecnica o arbitral del resto de funciones del club, manteniendo un reparto de responsabilidades claro dentro de la aplicacion.", None),
        ("INTRODUCIR CAPTURAS DE LA ZONA DE ARBITRO Y DE LA GESTION DE ACTAS", None),
        ("9.7 Cierre de sesion y recomendaciones de uso", "Heading 2"),
        ("Cuando el usuario termina de trabajar con la aplicacion, lo recomendable es cerrar la sesion para evitar accesos no deseados, especialmente si se esta utilizando un equipo compartido. Del mismo modo, conviene revisar bien la informacion antes de guardar cambios importantes o eliminar elementos del sistema.", None),
        ("En el caso concreto del gerente, resulta especialmente aconsejable revisar periodicamente las solicitudes de contacto pendientes para que la web no se convierta solo en un escaparate, sino tambien en un canal de atencion realmente util para las personas que contactan con el club.", None),
    ]


def build_section_10() -> list[tuple[str, str | None]]:
    return [
        ("El manual tecnico recoge la informacion necesaria para comprender la estructura interna del proyecto, sus principales componentes y la forma en la que se coordinan entre si. Este apartado esta pensado sobre todo para facilitar el mantenimiento del sistema y para dejar constancia de las decisiones tecnicas adoptadas durante el desarrollo.", None),
        ("10.1 Estructura general del sistema", "Heading 2"),
        ("BASKET_ALJARAFE sigue una arquitectura separada en tres grandes bloques: frontend, backend y base de datos. El frontend se encarga de la interfaz y de la interaccion con el usuario; el backend centraliza la logica de negocio, la seguridad y el acceso a datos; y MySQL actua como sistema persistente para almacenar la informacion necesaria del proyecto.", None),
        ("El despliegue mediante Docker refuerza esta separacion, ya que cada una de estas piezas puede ejecutarse dentro de su propio contenedor, facilitando tanto la organizacion del entorno como la portabilidad del sistema a otros equipos o servidores.", None),
        ("10.2 Estructura del frontend", "Heading 2"),
        ("El frontend del proyecto ha sido desarrollado con Angular. La aplicacion cliente se organiza a partir de componentes, rutas y servicios, permitiendo dividir la interfaz en bloques reutilizables y mantener una estructura relativamente clara. Existen componentes destinados a la zona publica, al inicio de sesion y a las distintas areas privadas segun el rol del usuario autenticado.", None),
        ("Entre estas pantallas se encuentra ahora tambien la pagina publica de contacto, asi como la vista privada desde la que el gerente puede consultar y eliminar las solicitudes recibidas. Esta ampliacion sigue el mismo esquema del resto del frontend: un componente dedicado para la interfaz, rutas para la navegacion y llamadas a servicios encargados de comunicarse con el backend.", None),
        ("Desde un punto de vista tecnico, Angular se apoya en servicios para realizar peticiones HTTP al servidor. Esto permite concentrar la comunicacion con la API en clases especificas y dejar los componentes centrados en la presentacion y gestion del estado de la vista.", None),
        ("INTRODUCIR CAPTURAS DE LA ESTRUCTURA DEL FRONTEND Y DE LAS RUTAS PRINCIPALES", None),
        ("10.3 Estructura del backend", "Heading 2"),
        ("El backend ha sido desarrollado con Spring Boot y sigue una organizacion basada en capas. Por una parte se encuentran los controladores, que reciben las peticiones HTTP. Por otra parte aparecen los servicios, donde se concentra la logica de negocio. Finalmente, el acceso a la base de datos se articula mediante entidades y repositorios.", None),
        ("Esta separacion ayuda a que el codigo resulte mas mantenible, ya que cada bloque tiene una responsabilidad concreta. Los controladores no deberian contener logica compleja, los servicios actuan como punto central de coordinacion y las entidades representan la informacion persistida en la base de datos.", None),
        ("En esta parte del backend se han añadido endpoints especificos para la nueva funcionalidad de solicitudes de contacto. Por un lado, la parte publica permite registrar una nueva solicitud enviada desde la web. Por otro lado, la parte privada del gerente permite listar las solicitudes existentes y eliminar aquellas que ya hayan sido gestionadas.", None),
        ("10.4 Seguridad y control de acceso", "Heading 2"),
        ("La aplicacion utiliza un sistema de autenticacion y autorizacion que diferencia claramente los accesos publicos de los privados. Segun la configuracion implementada, cada usuario autenticado accede solo a las rutas y funcionalidades acordes a su rol. Esto resulta esencial para evitar que un perfil sin permisos pueda acceder a operaciones administrativas o a informacion que no le corresponde.", None),
        ("Ademas, el backend protege sus endpoints privados para que solo puedan ser utilizados por usuarios con sesion valida. De esta forma, aunque alguien conociese una ruta interna, no podria operar con ella si no cumple las condiciones de acceso establecidas por la aplicacion.", None),
        ("10.5 Base de datos y persistencia", "Heading 2"),
        ("La base de datos del proyecto se ha modelado en MySQL y recoge tanto la informacion principal del club como los datos auxiliares necesarios para el funcionamiento de la aplicacion. Las entidades del backend representan este modelo de datos y se relacionan entre si a traves de JPA e Hibernate.", None),
        ("La nueva funcionalidad de contacto introduce tambien el almacenamiento persistente de las solicitudes enviadas desde la pagina publica. Esto permite que los mensajes no dependan de memoria temporal ni de elementos externos, sino que queden registrados en el propio sistema hasta que el gerente decida eliminarlos tras su gestion.", None),
        ("A nivel de despliegue, la persistencia se garantiza mediante un volumen asociado al contenedor de MySQL, lo que permite conservar la informacion aunque los contenedores del proyecto se detengan o se vuelvan a levantar mas adelante.", None),
        ("10.6 Scripts, inicializacion y datos de ejemplo", "Heading 2"),
        ("El proyecto incorpora scripts de inicializacion de base de datos y carga de datos de ejemplo para facilitar el arranque. Esto es especialmente util durante la fase de pruebas, demostracion y desarrollo, ya que permite levantar el sistema con una base ya preparada sin necesidad de introducir toda la informacion manualmente desde cero.", None),
        ("Desde el punto de vista del mantenimiento, este tipo de scripts tambien ayuda a reproducir entornos con mayor facilidad, algo importante cuando se quiere mover la aplicacion de un equipo local a otro o llevarla a un servidor remoto.", None),
        ("10.7 Despliegue tecnico y uso de Docker Hub", "Heading 2"),
        ("A nivel tecnico, el despliegue puede realizarse de dos formas principales. La primera consiste en construir las imagenes localmente a partir de los Dockerfile del proyecto. La segunda consiste en reutilizar imagenes ya publicadas en Docker Hub, lo que permite reducir trabajo en el servidor y hacer mas sencillo el proceso de instalacion o actualizacion.", None),
        ("Cuando se opta por Docker Hub, el servidor necesita basicamente el fichero compose, el fichero de variables de entorno y acceso a las imagenes publicadas. Esta estrategia resulta conveniente porque desacopla el servidor del proceso de compilacion y convierte el despliegue en una tarea mas cercana a la simple orquestacion de contenedores.", None),
        ("10.8 Mantenimiento y posibles ampliaciones", "Heading 2"),
        ("Desde el punto de vista tecnico, el proyecto queda preparado para seguir creciendo. La separacion por capas en el backend, la organizacion del frontend por componentes y la contenedorizacion del sistema hacen que sea viable incorporar nuevas funcionalidades en el futuro sin necesidad de rehacer por completo la estructura existente.", None),
        ("Como posibles ampliaciones, se podria profundizar en la monitorizacion, en la gestion de logs, en una estrategia mas avanzada de copias de seguridad o en un sistema de despliegue continuo apoyado en actualizaciones automatizadas de imagenes publicadas en Docker Hub.", None),
        ("INTRODUCIR CAPTURAS DE LA ESTRUCTURA INTERNA DEL PROYECTO Y DE LOS ELEMENTOS TECNICOS MAS RELEVANTES", None),
    ]


def update_document() -> None:
    document = Document(TARGET_DOC)

    zone_public_anchor = find_paragraph(document, "6.1 Zona pública")
    append_block(zone_public_anchor, build_contact_public_block())

    gerente_anchor = find_paragraph(document, "6.3 Gestión de gerente")
    append_block(gerente_anchor, build_contact_admin_block())

    section_8_anchor = clear_section(document, "8. Despliegue e instalación", "9. Manual de usuario")
    append_block(section_8_anchor, build_section_8())

    section_9_anchor = clear_section(document, "9. Manual de usuario", "10. Manual técnico")
    append_block(section_9_anchor, build_section_9())

    section_10_anchor = clear_section(document, "10. Manual técnico", "11. Pruebas realizadas y mejoras pendientes")
    append_block(section_10_anchor, build_section_10())

    document.save(TARGET_DOC)


if __name__ == "__main__":
    update_document()
